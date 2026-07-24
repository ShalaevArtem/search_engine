import os
import re
import sys
from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QTextBrowser
from PyQt6.QtCore import Qt

class PreviewDialog(QDialog):
    def __init__(self, file_path: str, highlight_terms: list, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Предпросмотр: {os.path.basename(file_path)}")
        self.setMinimumSize(800, 600)
        self._apply_dark_theme()

        self.terms = []
        if isinstance(highlight_terms, list):
            for t in highlight_terms:
                if isinstance(t, bytes):
                    try:
                        self.terms.append(t.decode('utf-8'))
                    except:
                        pass
                elif isinstance(t, str) and len(t) > 1:
                    self.terms.append(t)
        if not self.terms and hasattr(parent, 'last_query') and parent.last_query:
            self.terms = [parent.last_query]

        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(16, 16, 16, 16)

        terms_str = ", ".join(self.terms[:5]) if self.terms else "поиск по дате/имени"
        info = QLabel(f"Совпавшие термы: <b>{terms_str}</b>")
        info.setWordWrap(True)
        layout.addWidget(info)

        self.browser = QTextBrowser()
        self.browser.setReadOnly(True)
        layout.addWidget(self.browser)

        btns = QHBoxLayout()
        btn_open = QPushButton("Открыть в системе")
        btn_close = QPushButton("Закрыть")
        btn_open.clicked.connect(lambda: (os.startfile(file_path) if sys.platform == "win32" else None, self.close()))
        btn_close.clicked.connect(self.close)
        btns.addWidget(btn_open)
        btns.addStretch()
        btns.addWidget(btn_close)
        layout.addLayout(btns)

        self.load_and_highlight(file_path)

    def _apply_dark_theme(self):
        self.setStyleSheet("""
            QDialog {
                background-color: #1E1E1E;
                color: #D4D4D4;
            }
            QLabel {
                color: #D4D4D4;
                font-size: 13px;
            }
            QTextBrowser {
                background-color: #252526;
                color: #D4D4D4;
                border: 1px solid #3E3E42;
                border-radius: 4px;
                font-size: 13px;
                padding: 8px;
            }
            QPushButton {
                background-color: #0E639C;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-size: 13px;
                min-height: 32px;
            }
            QPushButton:hover {
                background-color: #1177BB;
            }
            QPushButton:pressed {
                background-color: #094771;
            }
        """)

    def load_and_highlight(self, path: str):
        try:
            ext = os.path.splitext(path)[1].lower()
            text = ""
            MAX_CHARS = 50000
            if ext == ".txt":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read(MAX_CHARS)
            elif ext == ".pdf":
                import pypdfium2 as pdfium
                doc = pdfium.PdfDocument(path)
                for i in range(min(len(doc), 25)):
                    page = doc[i]
                    text_page = page.get_textpage()
                    t = text_page.get_text_bounded()
                    if t:
                        text += f"\n[Стр. {i + 1}]\n" + t
                    if len(text) > MAX_CHARS:
                        break
            elif ext == ".docx":
                from docx import Document
                for p in Document(path).paragraphs:
                    if p.text.strip():
                        text += p.text + "\n"
                    if len(text) > MAX_CHARS:
                        break
            if not text.strip():
                self.browser.setHtml("<h3 style='color:#D4D4D4'>Текст не извлечён</h3>")
                return
            if not self.terms:
                self.browser.setPlainText(text[:MAX_CHARS])
                return
            safe_terms = [re.escape(t) for t in self.terms if t]
            if not safe_terms:
                self.browser.setPlainText(text[:MAX_CHARS])
                return
            pattern = re.compile('|'.join(safe_terms[:10]), re.IGNORECASE)
            safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            highlighted = pattern.sub(
                lambda
                    m: f"<mark style='background:#FFD700; color:#000; padding:2px 4px; border-radius:3px;'>{m.group(0)}</mark>",
                safe_text
            )
            self.browser.setHtml(
                f"<div style='font-family:Consolas, monospace; line-height:1.5; color:#D4D4D4;'>{highlighted}</div>")
        except Exception as e:
            import traceback
            self.browser.setHtml(f"""
                <h3 style='color:#F85149'>Ошибка предпросмотра</h3>
                <pre style='background:#252526; color:#D4D4D4; padding:8px; border-radius:4px; white-space:pre-wrap;'>{traceback.format_exc()[:500]}</pre>
            """)