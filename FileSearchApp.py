import subprocess
import sys
import os
import pathlib
import re
import time

# --- Исправление путей к ресурсам при запуске из .exe (PyInstaller) ---
if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
    # Для pymorphy2 словарей
    os.environ["PYMORPHY2_DICT_PATH"] = str(pathlib.Path(sys._MEIPASS).joinpath('pymorphy2_dicts_ru/data'))
if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
    # Для ru-synonyms
    os.environ["PYCHARTS_GALLERY_API"] = str(pathlib.Path(sys._MEIPASS).joinpath('ru-synonyms'))

# --- Импорт библиотек для работы с русским языком ---
import ru_synonyms
import pymorphy2_dicts_ru

import logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

import sys
from pathlib import Path
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QDateEdit, QTableWidget, QTableWidgetItem,
    QFileDialog, QMessageBox, QHeaderView, QProgressBar, QDialog, QTextEdit, QFormLayout, QInputDialog, QCheckBox,
    QTextBrowser
)
from PyQt6.QtCore import Qt, QDate, QThread, pyqtSignal

# --- Импорт модулей поиска и индексации ---
from core.indexer import FileIndexer
from core.searcher import (
    setup_search_parser, search_index, search_time_range,
    combined_search, search_by_filename, validate_query
)
from core.auth_manager import auth_manager
from core.access_control import filter_results_by_access, check_file_access
from config import Config

class IndexThread(QThread):
    """Отдельный поток для индексации файлов с поддержкой прогресса."""
    progress = pyqtSignal(int)
    finished = pyqtSignal(int, int, str)  # success, failed, message

    def __init__(self, directory: Path):
        super().__init__()
        self.directory = directory

    def run(self):
        def progress_callback(value):
            self.progress.emit(value)
        try:
            success, failed = FileIndexer.index_files(self.directory, progress_callback)
            self.finished.emit(success, failed, "Индексация завершена")
        except Exception as e:
            self.finished.emit(0, 0, f"Ошибка: {str(e)}")

class LoginDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Вход в систему")
        self.setMinimumSize(380, 220)
        self.setStyleSheet("""
            QDialog { background-color: #18191c; color: #e0e0e0; font-size: 12px; }
            QLineEdit { background-color: #2c2f33; border: 1px solid #444; padding: 6px; color: #e0e0e0; border-radius: 4px; }
            QPushButton { background-color: #7289da; color: white; border: none; padding: 8px 16px; border-radius: 4px; font-weight: bold; }
            QPushButton:hover { background-color: #5f73bc; }
            QLabel.error { color: #ff5555; font-size: 11px; min-height: 16px; }
            QCheckBox { color: #aaa; }
        """)

        layout = QFormLayout(self)
        layout.setSpacing(10)
        layout.setContentsMargins(20, 20, 20, 20)

        self.user_input = QLineEdit()
        self.user_input.setPlaceholderText("Логин")
        self.pass_input = QLineEdit()
        self.pass_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.pass_input.setPlaceholderText("Пароль")

        self.remember_cb = QCheckBox("Запомнить меня (30 дней)")
        self.error_label = QLabel()
        self.error_label.setObjectName("error")
        self.btn_login = QPushButton("Войти")

        layout.addRow("Логин:", self.user_input)
        layout.addRow("Пароль:", self.pass_input)
        layout.addRow(self.remember_cb)
        layout.addRow(self.error_label)
        layout.addRow(self.btn_login)

        self.btn_login.clicked.connect(self.try_login)
        self.pass_input.returnPressed.connect(self.try_login)

    def try_login(self):
        from core.auth_manager import auth_manager
        u, p = self.user_input.text().strip(), self.pass_input.text()
        if not u or not p:
            self.error_label.setText("Заполните все поля")
            return

        if auth_manager.create_session_token(u, p, self.remember_cb.isChecked()):
            self.done(QDialog.DialogCode.Accepted)
        else:
            self.error_label.setText("Неверный логин или пароль")
            self.pass_input.clear()

class UserManagementDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Управление пользователями")
        self.setMinimumSize(500, 350)
        self.setStyleSheet("""
            QDialog { background-color: #18191c; color: #e0e0e0; }
            QPushButton { background-color: #43b581; padding: 6px 12px; border-radius: 4px; }
            QPushButton:hover { background-color: #3ca374; }
            QPushButton.danger { background-color: #f04747; }
            QPushButton.danger:hover { background-color: #d63c3c; }
            QTableWidget { background-color: #23272a; color: #e0e0e0; gridline-color: #333; }
        """)

        layout = QVBoxLayout(self)
        self.table = QTableWidget(0, 4)
        self.table.setHorizontalHeaderLabels(["ID", "Логин", "Роль", "Активен"])
        layout.addWidget(self.table)

        btns = QHBoxLayout()
        self.btn_create = QPushButton("Создать")
        self.btn_toggle = QPushButton("Блок/Разблок")
        self.btn_delete = QPushButton("Удалить")
        self.btn_delete.setObjectName("danger")
        btns.addWidget(self.btn_create)
        btns.addWidget(self.btn_toggle)
        btns.addWidget(self.btn_delete)
        layout.addLayout(btns)

        self.btn_create.clicked.connect(self.create_user)
        self.btn_toggle.clicked.connect(self.toggle_active)
        self.btn_delete.clicked.connect(self.delete_user)

        self.load_users()

    def load_users(self):
        from models.database import SessionLocal, User, Role
        db = SessionLocal()
        try:
            users = db.query(User).all()
            self.table.setRowCount(len(users))
            for i, u in enumerate(users):
                self.table.setItem(i, 0, QTableWidgetItem(str(u.id)))
                self.table.setItem(i, 1, QTableWidgetItem(u.username))
                roles = ", ".join([r.name for r in u.roles])
                self.table.setItem(i, 2, QTableWidgetItem(roles))
                self.table.setItem(i, 3, QTableWidgetItem("✅" if u.is_active else "❌"))
        finally:
            db.close()

    def create_user(self):
        username, ok1 = QInputDialog.getText(self, "Новый пользователь", "Логин:")
        if not ok1 or not username.strip(): return
        password, ok2 = QInputDialog.getText(self, "Новый пользователь", "Пароль:", QLineEdit.EchoMode.Password)
        if not ok2: return

        from models.database import SessionLocal, User, Role
        from core.auth_manager import hash_password
        db = SessionLocal()
        try:
            if db.query(User).filter(User.username == username).first():
                QMessageBox.warning(self, "Ошибка", "Пользователь уже существует")
                return
            role = db.query(Role).filter(Role.name == "user").first()
            new_user = User(username=username.strip(), password_hash=hash_password(password), is_active=True)
            new_user.roles.append(role)
            db.add(new_user)
            db.commit()
            self.load_users()
        finally:
            db.close()

    def toggle_active(self):
        row = self.table.currentRow()
        if row == -1: return
        uid = int(self.table.item(row, 0).text())
        from models.database import SessionLocal, User
        db = SessionLocal()
        try:
            u = db.query(User).filter(User.id == uid).first()
            if u:
                u.is_active = not u.is_active
                db.commit()
                self.load_users()
        finally:
            db.close()

    def delete_user(self):
        row = self.table.currentRow()
        if row == -1: return
        uid = int(self.table.item(row, 0).text())
        if QMessageBox.question(self, "Подтверждение", "Удалить пользователя?") != QMessageBox.StandardButton.Yes: return
        from models.database import SessionLocal, User
        db = SessionLocal()
        try:
            db.query(User).filter(User.id == uid).delete()
            db.commit()
            self.load_users()
        finally:
            db.close()

class ACLManagementDialog(QDialog):
    """Диалог управления правилами доступа к файлам/папкам."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("📂 Управление доступом (ACL)")
        self.setMinimumSize(650, 400)
        self.setStyleSheet("""
            QDialog { background-color: #18191c; color: #e0e0e0; }
            QPushButton { background-color: #7289da; padding: 6px 12px; border-radius: 4px; }
            QPushButton:hover { background-color: #5f73bc; }
            QPushButton.danger { background-color: #f04747; }
            QTableWidget { background-color: #23272a; color: #e0e0e0; gridline-color: #333; }
        """)

        layout = QVBoxLayout(self)
        self.table = QTableWidget(0, 4)
        self.table.setHorizontalHeaderLabels(["ID", "Маска пути", "Доступно ролям", "Рекурсивно"])
        self.table.horizontalHeader().setStretchLastSection(True)
        layout.addWidget(self.table)

        btns = QHBoxLayout()
        self.btn_add = QPushButton("Добавить правило")
        self.btn_delete = QPushButton("Удалить выбранное")
        self.btn_delete.setObjectName("danger")
        btns.addWidget(self.btn_add)
        btns.addWidget(self.btn_delete)
        layout.addLayout(btns)

        self.btn_add.clicked.connect(self.add_rule)
        self.btn_delete.clicked.connect(self.delete_rule)
        self.load_rules()

    def load_rules(self):
        from models.database import SessionLocal, DocumentACL
        db = SessionLocal()
        try:
            rules = db.query(DocumentACL).all()
            self.table.setRowCount(len(rules))
            for i, r in enumerate(rules):
                self.table.setItem(i, 0, QTableWidgetItem(str(r.id)))
                self.table.setItem(i, 1, QTableWidgetItem(r.path_mask))
                self.table.setItem(i, 2, QTableWidgetItem(r.allowed_roles))
                self.table.setItem(i, 3, QTableWidgetItem("Да" if getattr(r, 'is_recursive', False) else "Нет"))
        finally:
            db.close()

    def add_rule(self):
        path, ok1 = QInputDialog.getText(self, "Новое правило", "Маска пути (напр. D:\\Docs\\* или *.pdf):")
        if not ok1 or not path.strip(): return
        roles, ok2 = QInputDialog.getText(self, "Новое правило", "Роли через запятую (admin,user):")
        if not ok2 or not roles.strip(): return
        recursive = QMessageBox.question(self, "Рекурсия?",
                                         "Применять к вложенным файлам и папкам?") == QMessageBox.StandardButton.Yes

        from models.database import SessionLocal, DocumentACL
        db = SessionLocal()
        try:
            if db.query(DocumentACL).filter(DocumentACL.path_mask == path.strip()).first():
                QMessageBox.warning(self, "Ошибка", "Такое правило уже существует")
                return
            db.add(DocumentACL(
                path_mask=path.strip().lower(),
                allowed_roles=roles.strip().replace(" ", ""),
                is_recursive=recursive
            ))
            db.commit()
            self.load_rules()
            # Сбрасываем кэш ACL, чтобы новые правила применились сразу
            from core.access_control import get_acl_rules
            import core.access_control
            core.access_control._cached_acl = None
        finally:
            db.close()

    def delete_rule(self):
        row = self.table.currentRow()
        if row == -1: return
        rule_id = int(self.table.item(row, 0).text())
        mask = self.table.item(row, 1).text()
        if QMessageBox.question(self, "Подтверждение",
                                f"Удалить правило '{mask}'?") != QMessageBox.StandardButton.Yes: return

        from models.database import SessionLocal, DocumentACL
        db = SessionLocal()
        try:
            db.query(DocumentACL).filter(DocumentACL.id == rule_id).delete()
            db.commit()
            self.load_rules()
            from core.access_control import get_acl_rules
            import core.access_control
            core.access_control._cached_acl = None
        finally:
            db.close()

class FirstRunSetupDialog(QDialog):
    """Безопасное окно первой настройки"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Первая настройка системы")
        self.setMinimumSize(420, 210)
        self.setStyleSheet("QDialog { background-color: #18191c; color: #e0e0e0; }")

        layout = QFormLayout(self)
        self.user_in = QLineEdit()
        self.pass_in = QLineEdit()
        self.pass_in.setEchoMode(QLineEdit.EchoMode.Password)
        self.btn = QPushButton("Создать администратора")
        self.error_lbl = QLabel()
        self.error_lbl.setStyleSheet("color: #ff5555; font-size: 11px;")

        layout.addRow("Логин (≥3 символов):", self.user_in)
        layout.addRow("Пароль (≥6 символов):", self.pass_in)
        layout.addRow(self.error_lbl)
        layout.addRow(self.btn)

        self.btn.clicked.connect(self._save_admin)
        self.pass_in.returnPressed.connect(self._save_admin)

    def _save_admin(self):
        u, p = self.user_in.text().strip(), self.pass_in.text()
        if len(u) < 3 or len(p) < 6:
            self.error_lbl.setText("Логин ≥3, пароль ≥6 символов")
            return
        self.btn.setEnabled(False)
        self.error_lbl.setText("Создание учётной записи...")
        QApplication.processEvents()  # Обновляем UI перед тяжёлой операцией

        try:
            from models.database import SessionLocal, User, Role
            from core.auth_manager import hash_password
            db = SessionLocal()
            try:
                # Проверка на случай повторного запуска
                if db.query(User).first():
                    self.accept()
                    return

                admin_role = Role(name="admin", description="Полный доступ")
                user_role = Role(name="user", description="Стандартный доступ")
                admin_user = User(username=u, password_hash=hash_password(p), is_active=True)
                admin_user.roles.append(admin_role)
                db.add_all([admin_role, user_role, admin_user])
                db.commit()
                self.accept()
            finally:
                db.close()
        except Exception as e:
            self.btn.setEnabled(True)
            self.error_lbl.setText(f"Ошибка: {str(e)[:60]}")
            import traceback
            traceback.print_exc()  # Вывод в консоль для отладки

class FileSearchApp(QMainWindow):
    """Главный класс GUI поисковой системы."""
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Поисковая система")
        self.setMinimumSize(800, 500)

        # --- Индекс и парсер ---
        Config.INDEX_DIR.mkdir(exist_ok=True)
        self.ix = FileIndexer.get_index(Config.INDEX_DIR)
        if self.ix.is_empty():
            print("Индекс пуст. Выполните индексацию.")
        self.parser = setup_search_parser(self.ix.schema)

        # --- Основной layout ---
        central_widget = QWidget()
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(8, 6, 8, 6)
        main_layout.setSpacing(6)

        # --- Прогресс-бар индексации (по умолчанию скрыт) ---
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.hide()
        main_layout.addWidget(self.progress_bar)

        # --- Вкладки для разных видов поиска и индексации ---
        self.tabs = QTabWidget()
        self.tabs.addTab(self.create_index_tab(), "📁 Индексация")
        self.tabs.addTab(self.create_keywords_tab(), "🔍 Ключевые слова")
        self.tabs.addTab(self.create_date_tab(), "📅 Дата")
        self.tabs.addTab(self.create_combined_tab(), "⚡ Комбинированный")
        self.tabs.addTab(self.create_filename_tab(), "📝 Имя файла")
        main_layout.addWidget(self.tabs)

        # --- Таблица для отображения результатов поиска ---
        self.results_table = QTableWidget(0, 3)
        self.results_table.setHorizontalHeaderLabels(["Имя", "Релевантность", "Дата изменения"])
        self.results_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.results_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.results_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.results_table.setAlternatingRowColors(True)
        self.results_table.setStyleSheet("font-size: 11px; padding: 2px;")
        self.results_table.hide()
        main_layout.addWidget(self.results_table)
        self.results_table.cellDoubleClicked.connect(self.open_file_from_result)

        self.setCentralWidget(central_widget)
        self.setStyleSheet(self.dark_stylesheet())

        self.index_thread = None
        self.last_query = ""

    def dark_stylesheet(self):
        """CSS-стили для тёмной темы интерфейса."""
        return """
        QMainWindow, QWidget {
            background-color: #18191c;
            color: #e0e0e0;
            font-family: 'Segoe UI', 'Arial', sans-serif;
            font-size: 11px;
        }
        QTabWidget::pane {
            border: 1px solid #23272a;
            background-color: #23272a;
            border-radius: 6px;
        }
        QTabBar::tab {
            background-color: #23272a;
            padding: 4px 10px;
            margin: 1px;
            border-top-left-radius: 4px;
            border-top-right-radius: 4px;
            font-size: 11px;
        }
        QTabBar::tab:selected {
            background-color: #2c2f33;
            color: #7289da;
            border-bottom: 2px solid #7289da;
        }
        QLineEdit, QDateEdit {
            background-color: #2c2f33;
            border: 1px solid #333;
            padding: 2px;
            color: #e0e0e0;
            border-radius: 3px;
            font-size: 11px;
        }
        QPushButton {
            background-color: #7289da;
            color: white;
            border: none;
            padding: 4px 10px;
            font-size: 11px;
            border-radius: 3px;
        }
        QPushButton:hover {
            background-color: #5f73bc;
        }
        QTableWidget {
            background-color: #23272a;
            alternate-background-color: #252525;
            gridline-color: #333;
            color: #e0e0e0;
            font-size: 11px;
        }
        QTableWidget::item {
            padding: 1px;
        }
        QHeaderView::section {
            background-color: #2c2f33;
            padding: 4px;
            border: 1px solid #333;
            border-radius: 3px;
            font-size: 11px;
        }
        QProgressBar {
            border: 1px solid #333;
            border-radius: 4px;
            text-align: center;
            color: #e0e0e0;
            background-color: #2c2f33;
            height: 14px;
            font-size: 11px;
        }
        QProgressBar::chunk {
            background-color: #7289da;
            border-radius: 4px;
        }
        """

    # --- Вкладка "Индексация" ---
    def create_index_tab(self):
        tab = QWidget()
        layout = QHBoxLayout(tab)
        layout.setSpacing(4)
        layout.addWidget(QLabel("Путь к папке:"))
        self.dir_input = QLineEdit()
        layout.addWidget(self.dir_input)
        browse_btn = QPushButton("Обзор")
        browse_btn.clicked.connect(self.browse_directory)
        layout.addWidget(browse_btn)
        index_btn = QPushButton("Индексировать")
        index_btn.clicked.connect(self.start_indexing)
        layout.addWidget(index_btn)
        return tab

    # --- Вкладка "Ключевые слова" ---
    def create_keywords_tab(self):
        tab = QWidget()
        layout = QHBoxLayout(tab)
        layout.setSpacing(4)
        layout.addWidget(QLabel("Запрос:"))
        self.keywords_input = QLineEdit()
        self.keywords_input.setPlaceholderText("Например: отчёт по проекту")
        layout.addWidget(self.keywords_input)
        search_btn = QPushButton("Поиск")
        search_btn.clicked.connect(self.search_keywords)
        layout.addWidget(search_btn)
        return tab

    # --- Вкладка "Дата" ---
    def create_date_tab(self):
        tab = QWidget()
        layout = QHBoxLayout(tab)
        layout.setSpacing(4)
        layout.addWidget(QLabel("Начальная дата:"))
        self.start_date = QDateEdit()
        self.start_date.setCalendarPopup(True)
        self.start_date.setDate(QDate.currentDate().addDays(-30))
        layout.addWidget(self.start_date)
        layout.addWidget(QLabel("Конечная дата:"))
        self.end_date = QDateEdit()
        self.end_date.setCalendarPopup(True)
        self.end_date.setDate(QDate.currentDate())
        layout.addWidget(self.end_date)
        search_btn = QPushButton("Поиск")
        search_btn.clicked.connect(self.search_date)
        layout.addWidget(search_btn)
        return tab

    # --- Вкладка "Комбинированный поиск" ---
    def create_combined_tab(self):
        tab = QWidget()
        vlayout = QVBoxLayout(tab)
        vlayout.setSpacing(3)
        hlayout1 = QHBoxLayout()
        hlayout1.addWidget(QLabel("Запрос:"))
        self.combined_query = QLineEdit()
        self.combined_query.setPlaceholderText("Ключевые слова...")
        hlayout1.addWidget(self.combined_query)
        vlayout.addLayout(hlayout1)
        hlayout2 = QHBoxLayout()
        hlayout2.addWidget(QLabel("Начальная дата:"))
        self.combined_start_date = QDateEdit()
        self.combined_start_date.setCalendarPopup(True)
        self.combined_start_date.setDate(QDate.currentDate().addDays(-30))
        hlayout2.addWidget(self.combined_start_date)
        hlayout2.addWidget(QLabel("Конечная дата:"))
        self.combined_end_date = QDateEdit()
        self.combined_end_date.setCalendarPopup(True)
        self.combined_end_date.setDate(QDate.currentDate())
        hlayout2.addWidget(self.combined_end_date)
        vlayout.addLayout(hlayout2)
        search_btn = QPushButton("Поиск")
        search_btn.clicked.connect(self.search_combined)
        vlayout.addWidget(search_btn, alignment=Qt.AlignmentFlag.AlignLeft)
        return tab

    # --- Вкладка "Имя файла" ---
    def create_filename_tab(self):
        tab = QWidget()
        layout = QHBoxLayout(tab)
        layout.setSpacing(4)
        layout.addWidget(QLabel("Имя файла:"))
        self.filename_input = QLineEdit()
        self.filename_input.setPlaceholderText("Например: report.docx")
        layout.addWidget(self.filename_input)
        search_btn = QPushButton("Поиск")
        search_btn.clicked.connect(self.search_filename)
        layout.addWidget(search_btn)
        return tab

    # --- Действия пользователя ---
    def browse_directory(self):
        """Открывает диалог выбора директории и записывает путь в поле."""
        directory = QFileDialog.getExistingDirectory(self, "Выберите директорию")
        if directory:
            self.dir_input.setText(directory)

    def start_indexing(self):
        """Запускает индексацию выбранной директории в отдельном потоке."""
        directory = self.dir_input.text()
        if not directory:
            QMessageBox.warning(self, "Внимание", "Укажите путь к директории!")
            return
        dir_path = Path(directory)
        if not dir_path.exists() or not dir_path.is_dir():
            QMessageBox.warning(self, "Ошибка", "Указанный путь не существует или не является директорией!")
            return

        self.progress_bar.setValue(0)
        self.progress_bar.show()
        self.tabs.setEnabled(False)

        self.index_thread = IndexThread(dir_path)
        self.index_thread.progress.connect(self.progress_bar.setValue)
        self.index_thread.finished.connect(self.on_indexing_finished)
        self.index_thread.start()

    def on_indexing_finished(self, success, failed, message):
        """Обработка завершения индексации: обновляет индекс и интерфейс."""
        self.progress_bar.hide()
        self.tabs.setEnabled(True)
        QMessageBox.information(self, "Индексация", f"{message}\nУспешно: {success}, Ошибок: {failed}")
        # Обновить индекс и парсер после индексации
        self.ix = FileIndexer.get_index(Config.INDEX_DIR)
        self.parser = setup_search_parser(self.ix.schema)

    def _get_user_roles(self):
        """Безопасно возвращает роли текущего пользователя или пустой список."""
        from core.auth_manager import auth_manager
        user = auth_manager.get_current_user()
        return user.roles if user and user.roles else []

    def display_results(self, results):
        """Отображает результаты + защита от Side-Channel (утечки времени)"""
        start_time = time.time()  # Фиксируем время начала

        self.results_table.setRowCount(0)

        self.results_table.setColumnCount(4)
        self.results_table.setHorizontalHeaderLabels(["Путь", "Релевантность", "Дата", "Действие"])
        self.results_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.results_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)

        if not results:
            self.results_table.hide()
            QMessageBox.information(self, "Поиск", "Ничего не найдено или нет прав доступа.")
            # Имитация задержки
            time.sleep(0.3)
            return

        try:
            for i, res in enumerate(results):
                self.results_table.insertRow(i)

                path = str(res.get('path', ''))
                score = float(res.get('score', 0.0))
                date = str(res.get('last_modified', ''))

                self.results_table.setItem(i, 0, QTableWidgetItem(path))
                self.results_table.setItem(i, 1, QTableWidgetItem(f"{score:.2f}"))
                self.results_table.setItem(i, 2, QTableWidgetItem(date))

                btn = QPushButton("👁")
                btn.setFixedWidth(32)
                btn.setStyleSheet("QPushButton { background-color: #7289da; color: white; border-radius: 3px; }")

                terms = res.get('matched_terms', [])
                btn.clicked.connect(
                    lambda checked=False, p=path, t=terms: PreviewDialog(p, t, self).exec()
                )

                self.results_table.setCellWidget(i, 3, btn)

            self.results_table.show()

            # Если поиск был слишком быстрым, добавляем задержку до 300мс
            elapsed = time.time() - start_time
            if elapsed < 0.3:
                time.sleep(0.3 - elapsed)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка отображения", f"Не удалось показать результаты:\n{e}")
            self.results_table.hide()

    def _open_preview(self, file_path: str):
        """Обёртка для открытия предпросмотра (изолирует логику от таблицы)."""
        if not hasattr(self, 'PreviewDialog'):
            # Если класс PreviewDialog ещё не добавлен, пока просто откроем файл системно
            self.open_file_from_result_by_path(file_path)
            return
        self.PreviewDialog(file_path, getattr(self, 'last_query', ''), self).exec()

    def search_keywords(self):
        self.last_query = self.keywords_input.text()
        if not self.last_query.strip():
            QMessageBox.warning(self, "Внимание", "Введите поисковый запрос!")
            return
        try:
            from core.searcher import search_index
            from core.access_control import filter_results_by_access
            with self.ix.searcher() as searcher:
                res = search_index(searcher, self.parser, self.last_query)
            self.display_results(filter_results_by_access(res, self._get_user_roles()))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка поиска", str(e))

    def search_date(self):
        try:
            from core.searcher import search_time_range
            from core.access_control import filter_results_by_access
            start = self.start_date.date().toString("yyyy-MM-dd")
            end = self.end_date.date().toString("yyyy-MM-dd")
            with self.ix.searcher() as searcher:
                res = search_time_range(searcher, start, end)
            self.display_results(filter_results_by_access(res, self._get_user_roles()))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка поиска", str(e))

    def search_combined(self):
        self.last_query = self.combined_query.text()
        if not self.last_query:
            QMessageBox.warning(self, "Внимание", "Введите запрос!")
            return
        try:
            from core.searcher import combined_search
            from core.access_control import filter_results_by_access
            params = {
                'query': self.last_query,
                'start_date': self.combined_start_date.date().toString("yyyy-MM-dd"),
                'end_date': self.combined_end_date.date().toString("yyyy-MM-dd"),
                'limit': 10
            }
            with self.ix.searcher() as searcher:
                res = combined_search(searcher, self.parser, params)
            self.display_results(filter_results_by_access(res, self._get_user_roles()))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка поиска", str(e))

    def search_filename(self):
        self.last_query = self.filename_input.text()
        if not self.last_query:
            QMessageBox.warning(self, "Внимание", "Введите имя файла!")
            return
        try:
            from core.searcher import search_by_filename
            from core.access_control import filter_results_by_access
            with self.ix.searcher() as searcher:
                res = search_by_filename(searcher, self.last_query)
            self.display_results(filter_results_by_access(res, self._get_user_roles()))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка поиска", str(e))

    def open_file_from_result(self, row, column):
        """Открывает файл из результатов поиска двойным кликом."""
        path_item = self.results_table.item(row, 0)
        if not path_item:
            return
        path = path_item.text()
        if not os.path.exists(path):
            QMessageBox.warning(self, "Ошибка", "Файл не найден!")
            return

        # Проверка прав перед открытием
        current_user = auth_manager.get_current_user()
        if not current_user or not check_file_access(path, current_user.roles):
            QMessageBox.warning(self, "Отказ в доступе", "У вас нет прав для открытия этого файла.")
            return

        try:
            if sys.platform == "win32":
                os.startfile(path)
            elif sys.platform == "darwin":
                import subprocess
                subprocess.run(["open", path])
            else:
                import subprocess
                subprocess.run(["xdg-open", path])
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть файл:\n{e}")

class PreviewDialog(QDialog):
    def __init__(self, file_path: str, highlight_terms: list, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Предпросмотр: {os.path.basename(file_path)}")
        self.setMinimumSize(700, 500)
        self.setStyleSheet("QDialog { background-color: #18191c; color: #e0e0e0; }")

        # Безопасное декодирование терминов
        self.terms = []
        if isinstance(highlight_terms, list):
            for t in highlight_terms:
                if isinstance(t, bytes):
                    try:
                        self.terms.append(t.decode('utf-8'))
                    except:
                        pass
                elif isinstance(t, str) and len(t) > 1:
                    self.terms.append(t)

        if not self.terms and hasattr(parent, 'last_query') and parent.last_query:
            self.terms = [parent.last_query]

        layout = QVBoxLayout(self)
        terms_str = ", ".join(self.terms[:5]) if self.terms else "поиск по дате/имени"
        info = QLabel(f"Совпавшие термы: <b>{terms_str}</b>")
        info.setWordWrap(True)
        layout.addWidget(info)

        self.browser = QTextBrowser()
        self.browser.setReadOnly(True)
        self.browser.setStyleSheet(
            "QTextBrowser { background-color: #23272a; color: #e0e0e0; font-size: 13px; padding: 8px; }")
        layout.addWidget(self.browser)

        btns = QHBoxLayout()
        btn_open = QPushButton("📂 Открыть в системе")
        btn_close = QPushButton("Закрыть")
        btn_open.clicked.connect(lambda: (os.startfile(file_path) if sys.platform == "win32" else None, self.close()))
        btn_close.clicked.connect(self.close)
        btns.addWidget(btn_open)
        btns.addStretch()
        btns.addWidget(btn_close)
        layout.addLayout(btns)

        self.load_and_highlight(file_path)

    def load_and_highlight(self, path: str):
        try:
            import os, re
            ext = os.path.splitext(path)[1].lower()
            text = ""
            MAX_CHARS = 50000

            if ext == ".txt":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read(MAX_CHARS)
            elif ext == ".pdf":
                import pypdfium2 as pdfium
                doc = pdfium.PdfDocument(path)
                for i in range(min(len(doc), 25)):
                    page = doc[i]
                    text_page = page.get_textpage()
                    t = text_page.get_text_bounded()
                    if t: text += f"\n📄 [Стр. {i + 1}]\n" + t
                    if len(text) > MAX_CHARS: break
            elif ext == ".docx":
                from docx import Document
                for p in Document(path).paragraphs:
                    if p.text.strip(): text += p.text + "\n"
                    if len(text) > MAX_CHARS: break

            if not text.strip():
                self.browser.setHtml("<h3>Текст не извлечён</h3>")
                return

            if not self.terms:
                self.browser.setPlainText(text[:MAX_CHARS])
                return

            # Экранирование и подсветка
            safe_terms = [re.escape(t) for t in self.terms if t]
            if not safe_terms:
                self.browser.setPlainText(text[:MAX_CHARS])
                return

            pattern = re.compile('|'.join(safe_terms[:10]), re.IGNORECASE)
            safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            highlighted = pattern.sub(
                lambda
                    m: f"<mark style='background:#f1c40f; color:#000; padding:2px 4px; border-radius:3px;'>{m.group(0)}</mark>",
                safe_text
            )
            self.browser.setHtml(f"<div style='font-family:Consolas, monospace; line-height:1.5;'>{highlighted}</div>")

        except Exception as e:
            import traceback
            self.browser.setHtml(f"""
                <h3 style='color:#f04747'>Ошибка предпросмотра</h3>
                <pre style='background:#2c2f33; padding:8px; border-radius:4px; white-space:pre-wrap;'>{traceback.format_exc()[:500]}</pre>
            """)

# --- Точка входа в приложение ---
if __name__ == "__main__":
    import sys
    from PyQt6.QtWidgets import QApplication, QDialog, QPushButton, QHBoxLayout
    from PyQt6.QtCore import QTimer
    from models.database import init_db, DB_PATH, engine
    from core.auth_manager import auth_manager
    import logging

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    logger = logging.getLogger(__name__)

    app = QApplication(sys.argv)
    # Держим процесс активным между переключениями окон
    app.setQuitOnLastWindowClosed(False)

    # Инициализация базы данных
    is_ready = init_db()
    if not is_ready and DB_PATH.exists():
        logger.warning("База данных существует, но пуста или повреждена. Пересоздаём...")
        engine.dispose()
        import time

        time.sleep(0.15)
        try:
            DB_PATH.unlink(missing_ok=True)
            for ext in ('-wal', '-shm'):
                p = DB_PATH.with_suffix(f'.db{ext}')
                if p.exists():
                    p.unlink()
        except PermissionError as e:
            logger.error(f"Не удалось удалить базу данных: {e}. Закройте IDE и удалите файл вручную.")
            sys.exit(1)
        is_ready = init_db()

    if not is_ready:
        setup_dlg = FirstRunSetupDialog()
        if setup_dlg.exec() != QDialog.DialogCode.Accepted:
            sys.exit(0)

    # Глобальная ссылка предотвращает удаление окна сборщиком мусора PyQt
    active_window = None

    def show_login():
        global active_window
        if auth_manager.check_stored_session():
            show_main_window()
            return

        active_window = LoginDialog()
        # Закрытие крестиком в окне входа завершает процесс
        active_window.rejected.connect(app.quit)
        # Успешный вход открывает главное окно
        active_window.accepted.connect(show_main_window)
        active_window.show()

    def show_main_window():
        global active_window
        active_window = FileSearchApp()

        # Перехват события закрытия для различения крестика и кнопки выхода
        def handle_close_event(event):
            if event.spontaneous():
                # Пользователь нажал системный крестик
                app.quit()
            event.accept()

        active_window.closeEvent = handle_close_event

        main_layout = active_window.centralWidget().layout()
        if main_layout:
            btns_layout = QHBoxLayout()
            btns_layout.setSpacing(8)

            if auth_manager.has_role("admin"):
                for btn_cls, label in [(UserManagementDialog, "Пользователи"), (ACLManagementDialog, "Доступ")]:
                    btn = QPushButton(label)
                    btn.setStyleSheet(
                        "QPushButton { background-color: #f04747; color: white; font-weight: bold; padding: 6px 12px; border-radius: 4px; }")
                    btn.clicked.connect(lambda checked=False, cls=btn_cls: cls(active_window).exec())
                    btns_layout.addWidget(btn)

            logout_btn = QPushButton("Выйти")
            logout_btn.setStyleSheet(
                "QPushButton { background-color: #43b581; color: white; font-weight: bold; padding: 6px 12px; border-radius: 4px; }")

            def on_logout():
                auth_manager.clear_session()
                active_window.close()
                # Небольшая задержка позволяет окну полностью закрыться перед показом нового
                QTimer.singleShot(100, show_login)

            logout_btn.clicked.connect(on_logout)
            btns_layout.addWidget(logout_btn)
            main_layout.addLayout(btns_layout)

        active_window.show()

    show_login()
    sys.exit(app.exec())