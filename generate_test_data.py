import os
import platform
import random
import string
from pathlib import Path
from docx import Document
from fpdf import FPDF  # pip install fpdf2

# ============================================================================
# НАСТРОЙКА КОЛИЧЕСТВА ФАЙЛОВ (меняй эти числа для разных серий тестов)
# ============================================================================
COUNT_BUH = 1200  # Бухгалтерия
COUNT_LAW = 1200  # Юристы
COUNT_COMMON = 1200  # Общие

# Настройка объёма текста внутри одного документа
PARAGRAPHS_PER_DOC = 50  # сколько параграфов в TXT и DOCX
PAGES_PER_PDF = 10  # условное количество "страниц" текста в PDF

BASE_DIR = Path("D:/TestDocs")  # <-- поменяй путь при необходимости

# ============================================================================
# СЛОВАРИ ТЕКСТОВ (шаблоны для генерации связных больших текстов)
# ============================================================================
SENTENCE_TEMPLATES = [
    "В соответствии с действующим законодательством {subject} обязан предоставить {object} в установленные сроки.",
    "Документ содержит сведения о {subject}, а также аналитические данные по {object}.",
    "На основании представленных материалов {subject} принял решение об утверждении {object}.",
    "В разделе {section} приведены показатели эффективности использования {object}.",
    "Ответственный исполнитель {subject} подготовил отчёт о выполнении работ по {object}.",
    "Во исполнение пункта {num} договора от {date} {subject} направляет {object}.",
    "Настоящим уведомляем, что {subject} проведена проверка состояния {object}.",
    "В целях повышения эффективности {subject} разработан план мероприятий по {object}.",
    "Согласно регламенту, {subject} должен согласовать {object} до наступления контрольной даты.",
    "В приложении №{num} содержится детальная информация о {subject} и связанных с ним {object}.",
]

SUBJECTS = [
    "общество с ограниченной ответственностью", "финансовый отдел", "главный бухгалтер",
    "юридическая служба", "отдел кадров", "совет директоров", "аудиторская комиссия",
    "налоговый инспектор", "контрагент", "поставщик", "покупатель", "акционер",
    "исполнительный директор", "комитет по закупкам", "ревизионная группа"
]

OBJECTS = [
    "бухгалтерская отчётность", "договорные документы", "кадровые приказы",
    "инвентаризационные описи", "налоговые декларации", "платёжные поручения",
    "акт выполненных работ", "протокол заседаний", "счета на оплату", "накладные",
    "справки о доходах", "трудовые договоры", "регламенты и инструкции", "лицензии"
]

SECTIONS = [
    "финансовый", "оперативный", "стратегический", "аналитический",
    "кадровый", "юридический", "налоговый", "административный"
]


def generate_sentence():
    """Генерирует одно осмысленное предложение."""
    tpl = random.choice(SENTENCE_TEMPLATES)
    return tpl.format(
        subject=random.choice(SUBJECTS),
        object=random.choice(OBJECTS),
        section=random.choice(SECTIONS),
        num=random.randint(1, 25),
        date=f"{random.randint(1, 28):02d}.{random.randint(1, 12):02d}.20{random.randint(20, 26)}"
    )


def generate_large_text(num_paragraphs: int) -> str:
    """Генерирует большой связный текст из N параграфов."""
    paragraphs = []
    for i in range(num_paragraphs):
        # 3-6 предложений в абзаце
        sentences = [generate_sentence() for _ in range(random.randint(3, 6))]
        paragraphs.append(" ".join(sentences))
    return "\n\n".join(paragraphs)


# ============================================================================
# СОЗДАНИЕ ФАЙЛОВ
# ============================================================================

def create_txt_file(filepath: Path, topic: str):
    text = f"Тема документа: {topic}.\n\n"
    text += generate_large_text(PARAGRAPHS_PER_DOC)
    text += f"\n\nДата формирования: {random.randint(1, 28):02d}.{random.randint(1, 12):02d}.2025."
    filepath.write_text(text, encoding="utf-8")
    print(f"  [TXT] {filepath} ({filepath.stat().st_size:,} байт)")


def create_docx_file(filepath: Path, topic: str):
    doc = Document()
    doc.add_heading(topic, level=1)

    # Добавляем большой текст частями, чтобы не перегружать память
    for _ in range(PARAGRAPHS_PER_DOC // 5):
        chunk = generate_large_text(5)
        for paragraph in chunk.split("\n\n"):
            doc.add_paragraph(paragraph)

    # Добавляем таблицу (иногда)
    if random.random() > 0.3:
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        for i in range(4):
            for j in range(3):
                table.cell(i, j).text = f"Данные {i + 1},{j + 1}"

    doc.save(filepath)
    print(f"  [DOCX] {filepath} ({filepath.stat().st_size:,} байт)")

def find_cyrillic_font() -> str:
    """Ищет системный TTF-шрифт с поддержкой кириллицы."""
    system = platform.system()
    candidates = []

    if system == "Windows":
        font_dir = Path("C:/Windows/Fonts")
        candidates = [
            font_dir / "arial.ttf",
            font_dir / "calibri.ttf",
            font_dir / "tahoma.ttf",
            font_dir / "times.ttf",
        ]
    elif system == "Linux":
        candidates = [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
            Path("/usr/share/fonts/truetype/freefont/FreeSans.ttf"),
        ]
    elif system == "Darwin":  # macOS
        candidates = [
            Path("/Library/Fonts/Arial.ttf"),
            Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
            Path("/usr/local/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ]

    # Проверяем текущую папку (если пользователь положил шрифт рядом)
    candidates += [Path("DejaVuSans.ttf"), Path("arial.ttf")]

    for path in candidates:
        if path.exists():
            return str(path)

    raise FileNotFoundError(
        "Не найден TTF-шрифт с кириллицей.\n"
        "Windows: используйте arial.ttf из C:\\Windows\\Fonts\n"
        "Linux: sudo apt install fonts-dejavu-core\n"
        "Или скачайте DejaVuSans.ttf и положите рядом со скриптом."
    )


def create_pdf_file(filepath: Path, topic: str):
    """PDF с нормальной кириллицей через системный шрифт."""
    font_path = find_cyrillic_font()

    pdf = FPDF()
    pdf.add_page()

    # Добавляем шрифт ТОЛЬКО в обычном начертании
    pdf.add_font("SysCyr", "", font_path, uni=True)

    # Заголовок — просто увеличиваем размер, без bold
    pdf.set_font("SysCyr", "", 16)
    pdf.cell(0, 10, topic, ln=True)
    pdf.ln(5)

    # Основной текст
    pdf.set_font("SysCyr", "", 11)

    # Генерируем большой текст
    text = generate_large_text(PARAGRAPHS_PER_DOC * 2)

    for paragraph in text.split("\n\n"):
        if pdf.get_y() > 270:
            pdf.add_page()
            pdf.set_font("SysCyr", "", 11)
        pdf.multi_cell(0, 6, paragraph)
        pdf.ln(2)

    pdf.output(str(filepath))
    print(f"  [PDF] {filepath} ({filepath.stat().st_size:,} байт)")


def generate_files(category: str, count: int, base_path: Path):
    category_path = base_path / category
    category_path.mkdir(parents=True, exist_ok=True)

    topics_pool = [
                      f"{category} — документ по договору №{i}" for i in range(1, count + 1)
                  ] + [
                      f"{category} — отчёт за {random.choice(['I', 'II', 'III', 'IV'])} квартал 2025"
                      for _ in range(count)
                  ]

    extensions = [".txt", ".pdf", ".docx"]

    for i in range(count):
        topic = random.choice(topics_pool)
        ext = random.choice(extensions)
        # Имя файла: латиница + номер, чтобы избежать проблем с длинными путями
        safe_topic = topic.replace(" ", "_").replace("—", "-")[:40]
        filename = f"{safe_topic}_{i + 1:03d}{ext}"
        filepath = category_path / filename

        try:
            if ext == ".txt":
                create_txt_file(filepath, topic)
            elif ext == ".docx":
                create_docx_file(filepath, topic)
            elif ext == ".pdf":
                create_pdf_file(filepath, topic)
        except Exception as e:
            print(f"  ОШИБКА создания {filepath}: {e}")


def main():
    print("=" * 60)
    print("Генератор тестовых документов (large-scale)")
    print("=" * 60)

    if BASE_DIR.exists():
        import shutil
        print(f"Удаление старой директории: {BASE_DIR}")
        shutil.rmtree(BASE_DIR)
    BASE_DIR.mkdir(parents=True, exist_ok=True)

    config = {
        "Бухгалтерия": COUNT_BUH,
        "Юристы": COUNT_LAW,
        "Общие": COUNT_COMMON,
    }

    total = 0
    total_size = 0
    for category, count in config.items():
        print(f"\nГенерация '{category}': {count} файлов...")
        generate_files(category, count, BASE_DIR)
        total += count

    # Подсчёт итогов
    print(f"\n{'=' * 60}")
    print(f"Готово! Создано {total} файлов.")
    for path in sorted(BASE_DIR.rglob("*")):
        if path.is_file():
            size = path.stat().st_size
            total_size += size
    print(f"Общий объём: {total_size / (1024 * 1024):.1f} МБ")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()